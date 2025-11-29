"""
Document Generator for Pic2Doc
Creates Word documents with images and captions
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Tuple, Dict, Any, Callable, Optional
from pathlib import Path
import math


class DocumentGenerator:
    """Generates Word documents with images and captions"""

    def __init__(self, config: Dict[str, Any]):
        """
        Initialize document generator

        Args:
            config: Configuration dictionary with formatting settings
        """
        self.config = config

    def _set_document_margins(self, doc: Document):
        """
        Set document margins from configuration

        Args:
            doc: Document object to modify
        """
        # Access document sections
        for section in doc.sections:
            # Convert cm to inches for docx (Word uses inches internally)
            section.top_margin = Cm(self.config.get('margin_top_cm', 1.27))
            section.bottom_margin = Cm(self.config.get('margin_bottom_cm', 1.27))
            section.left_margin = Cm(self.config.get('margin_left_cm', 1.27))
            section.right_margin = Cm(self.config.get('margin_right_cm', 1.27))

    def _calculate_optimal_grid(self, images_per_page: int) -> Tuple[int, int]:
        """
        Calculate optimal grid layout (cols x rows) for given number of images

        Args:
            images_per_page: Number of images to fit on one page

        Returns:
            Tuple of (cols, rows) for optimal space utilization
        """
        if images_per_page == 1:
            return (1, 1)
        elif images_per_page == 2:
            return (2, 1)  # 2 images side by side
        elif images_per_page == 3:
            return (2, 2)  # 2x2 grid, last cell empty
        elif images_per_page == 4:
            return (2, 2)  # Perfect 2x2 grid
        elif images_per_page == 5:
            return (3, 2)  # 3x2 grid, last cell empty
        elif images_per_page == 6:
            return (3, 2)  # Perfect 3x2 grid
        elif images_per_page == 7:
            return (3, 3)  # 3x3 grid, 2 cells empty
        elif images_per_page == 8:
            return (3, 3)  # 3x3 grid, 1 cell empty
        elif images_per_page == 9:
            return (3, 3)  # Perfect 3x3 grid
        elif images_per_page == 10:
            return (3, 4)  # 3x4 grid, 2 cells empty
        else:
            # For larger numbers, calculate optimal grid
            # Try to get close to square aspect ratio
            cols = math.ceil(math.sqrt(images_per_page))
            rows = math.ceil(images_per_page / cols)
            return (cols, rows)

    def _calculate_layout(self, images_per_page: int, page_images: List[Tuple]) -> List[List[int]]:
        """
        Calculate intelligent layout for images on a page to maximize space utilization
        Fills row by row in strict order from the image list

        Args:
            images_per_page: Target number of images per page
            page_images: List of image data tuples for this page

        Returns:
            List of rows, where each row contains indices of images in that row
            Example: [[0, 1], [2, 3]] = 2x2 grid layout (strict left-to-right, top-to-bottom order)
        """
        num_images = len(page_images)

        # Get optimal grid dimensions
        cols, rows = self._calculate_optimal_grid(images_per_page)

        # Build layout row by row in strict order
        layout = []
        img_idx = 0

        for row in range(rows):
            if img_idx >= num_images:
                break

            row_images = []
            for col in range(cols):
                if img_idx < num_images:
                    row_images.append(img_idx)
                    img_idx += 1
                else:
                    break

            if row_images:
                layout.append(row_images)

        return layout

    def _calculate_image_size(self, image_info: Optional[Dict], available_width: float,
                             num_in_row: int, images_per_page: int, total_rows: int,
                             font_size: int) -> Tuple[float, float]:
        """
        Calculate appropriate image size based on available space and grid layout
        Uses much more conservative estimates to prevent page breaks

        Args:
            image_info: Dict with image dimensions and orientation
            available_width: Available width in inches for all images in row
            num_in_row: Number of images in this row
            images_per_page: Total images expected per page
            total_rows: Total number of rows in grid
            font_size: Font size for captions in points

        Returns:
            Tuple of (width, height) in inches
        """
        if not image_info:
            # Fallback size
            width = available_width / num_in_row * 0.85
            return (width, width * 1.33)  # Assume portrait aspect ratio

        aspect_ratio = image_info.get('aspect_ratio', 1.0)

        # Calculate width per image with padding between images
        padding_between = 0.05  # Reduced padding
        total_padding = padding_between * (num_in_row - 1)
        width_per_image = (available_width - total_padding) / num_in_row

        # Calculate height based on aspect ratio
        height = width_per_image / aspect_ratio

        # VERY conservative height calculation to prevent ANY page breaks
        # A4 page is 11.7 inches tall
        usable_height = 8.5  # Very conservative estimate (was 9.0)

        # Calculate caption height - be more generous
        caption_height_per_row = (font_size / 72) * 2.5  # Increased from 2.0

        # More spacing between rows
        spacing_between_rows = 0.15  # Increased from 0.1

        # More table overhead
        table_overhead = 0.3  # Increased from 0.2

        # Calculate total space needed
        total_caption_space = caption_height_per_row * total_rows
        total_row_spacing = spacing_between_rows * max(0, total_rows - 1)

        # Available space for all images
        available_for_images = usable_height - total_caption_space - total_row_spacing - table_overhead

        # Max height per image
        max_height_per_image = available_for_images / total_rows

        # Reduce further for many columns
        if num_in_row > 2:
            max_height_per_image *= 0.90  # More conservative

        # Apply maximum height constraint
        if height > max_height_per_image:
            height = max_height_per_image
            width_per_image = height * aspect_ratio

        # Final safety reduction - reduce everything by 10%
        width_per_image *= 0.90
        height *= 0.90

        return (width_per_image, height)

    def _make_table_keep_together(self, table):
        """
        Apply keep-together properties to prevent table from breaking across pages

        Args:
            table: Table object to modify
        """
        # Set properties for the entire table to prevent page breaks
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Can't split table across pages
        cantSplit = OxmlElement('w:cantSplit')
        cantSplit.set(qn('w:val'), '1')
        tblPr.append(cantSplit)

        # Keep rows together
        for row in table.rows:
            trPr = row._element.get_or_add_trPr()
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '1')
            trPr.append(cantSplit)

    def create_document(
        self,
        image_data: List[Tuple[str, str, str, Optional[Dict]]],
        output_path: str,
        progress_callback: Optional[Callable[[int, int, str], None]] = None
    ) -> tuple[int, List[str]]:
        """
        Create Word document with images and captions using intelligent layout
        Maintains strict Excel sheet order

        Args:
            image_data: List of tuples (filename, caption, image_path, image_info_dict)
                       image_info_dict contains: orientation, width, height, aspect_ratio
                       Images are processed in the EXACT order they appear in this list
            output_path: Path where to save the document
            progress_callback: Optional callback function(current, total, filename)

        Returns:
            Tuple of (processed_count, error_list)
        """
        print("\nErstelle Word-Dokument...")

        doc = Document()
        self._set_document_margins(doc)

        # Set default font for document
        style = doc.styles['Normal']
        style.font.name = self.config['font_name']

        processed_count = 0
        missing_files = []
        error_details = []  # Store (filename, error_message) tuples
        total_images = len(image_data)
        images_per_page = self.config['images_per_page']

        # Calculate page width (A4 with margins)
        page_width_inches = 8.27 - (self.config.get('margin_left_cm', 1.27) + self.config.get('margin_right_cm', 1.27)) / 2.54

        # Process images in pages - STRICT ORDER from Excel
        for page_start in range(0, total_images, images_per_page):
            page_end = min(page_start + images_per_page, total_images)
            page_images = image_data[page_start:page_end]

            # Calculate layout for this page
            layout = self._calculate_layout(images_per_page, page_images)
            total_rows = len(layout)

            # Get max columns to create table
            max_cols = max(len(row) for row in layout)

            # Create one table for the entire page grid
            # Table has 2 rows per image row (image row + caption row)
            table = doc.add_table(rows=total_rows * 2, cols=max_cols)
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Make table stay together on one page
            self._make_table_keep_together(table)

            # Remove table borders
            for row in table.rows:
                for cell in row.cells:
                    cell_elem = cell._element
                    tc_pr = cell_elem.get_or_add_tcPr()
                    tc_borders = OxmlElement('w:tcBorders')
                    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                        border = OxmlElement(f'w:{border_name}')
                        border.set(qn('w:val'), 'none')
                        tc_borders.append(border)
                    tc_pr.append(tc_borders)

            # Fill the table with images IN STRICT ORDER
            for layout_row_idx, row_indices in enumerate(layout):
                num_in_row = len(row_indices)
                table_row_idx = layout_row_idx * 2  # Each layout row takes 2 table rows

                for col_idx, img_idx in enumerate(row_indices):
                    entry = page_images[img_idx]

                    if len(entry) == 4:
                        filename, caption, image_path, image_info = entry
                    else:
                        filename, caption, image_path = entry
                        image_info = None

                    try:
                        # Calculate size
                        img_width, img_height = self._calculate_image_size(
                            image_info,
                            page_width_inches,
                            num_in_row,
                            images_per_page,
                            total_rows,
                            self.config['font_size']
                        )

                        # Add image to table
                        cell = table.rows[table_row_idx].cells[col_idx]
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        run.add_picture(image_path, width=Inches(img_width))

                        # Add caption to next row
                        cell = table.rows[table_row_idx + 1].cells[col_idx]
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = paragraph.add_run(caption)

                        # Apply font formatting
                        caption_font = caption_run.font
                        caption_font.name = self.config['font_name']
                        caption_font.size = Pt(self.config['font_size'])
                        caption_font.bold = self.config['font_bold']
                        caption_font.italic = self.config['font_italic']
                        caption_font.underline = self.config['font_underline']

                        processed_count += 1
                        print(f"✓ Bild {processed_count}: {filename}")

                        if progress_callback:
                            progress_callback(processed_count, total_images, filename)

                    except FileNotFoundError as e:
                        error_msg = f"Datei nicht gefunden"
                        print(f"✗ Fehler bei {filename}: {error_msg}")
                        missing_files.append(filename)
                        error_details.append((filename, error_msg))
                    except Exception as e:
                        error_msg = str(e)
                        print(f"✗ Fehler bei {filename}: {error_msg}")
                        missing_files.append(filename)
                        error_details.append((filename, error_msg))

            # Add page break after each page (except last)
            if page_end < total_images:
                doc.add_page_break()

        # Save document
        output_path = Path(output_path)
        doc.save(str(output_path))

        # Print summary
        print(f"\n{'='*70}")
        print(f"✓ Word-Dokument erfolgreich erstellt!")
        print(f"  Gespeichert unter: {output_path}")
        print(f"  Bilder verarbeitet: {processed_count}/{total_images}")

        if error_details:
            print(f"\n⚠ {len(error_details)} Datei(en) mit Fehlern:")
            for fname, error_msg in error_details[:10]:
                print(f"  - {fname}")
                print(f"    Grund: {error_msg}")
            if len(error_details) > 10:
                print(f"  ... und {len(error_details) - 10} weitere")

        print(f"{'='*70}")

        return processed_count, missing_files
